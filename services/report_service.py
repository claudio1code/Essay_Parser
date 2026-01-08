from io import BytesIO
from typing import Any, Dict, Optional

from docx import Document
from docx.shared import Inches

from config import Config
from logger import get_logger

logger = get_logger(__name__)


def preencher_e_gerar_docx(
    dados: Dict[str, Any], caminho_template: str = Config.TEMPLATE_DOCX_PATH
) -> Optional[BytesIO]:
    """
    Preenche o template .docx preservando formatação e retorna o arquivo em bytes.

    Args:
        dados (Dict[str, Any]): Dicionário com os dados da correção.
        caminho_template (str): Caminho para o arquivo de template .docx.

    Returns:
        Optional[BytesIO]: Buffer contendo o documento gerado ou None em caso de erro.
    """
    try:
        logger.info(f"Gerando DOCX usando template: {caminho_template}")
        document = Document(caminho_template)
        comps = dados.get("analise_competencias", {})

        # Mapeamento de chaves para placeholders no template
        substituicoes = {
            "{{NOME_ALUNO}}": dados.get("nome_aluno", "Não informado"),
            "{{TEMA}}": dados.get("tema_redacao", "Não informado"),
            "{{ANO}}": dados.get("ano_turma", "Não informado"),
            "{{BIMESTRE}}": dados.get("bimestre", "Não informado"),
            "{{NOTA_FINAL}}": str(dados.get("nota_final", "N/A")),
            "{{COMENTARIOS}}": dados.get("comentarios_gerais", ""),
            "{{ALERTA_ORIGINALIDADE}}": dados.get("alerta_originalidade") or "",
        }

        # Adiciona as substituições para cada competência
        for i in range(1, 6):
            comp_data = comps.get(f"c{i}", {})
            nota = str(comp_data.get("nota", "N/A"))
            analise = comp_data.get("analise", "")

            # Limpa a análise removendo formatação específica que possa interferir (como Markdown)
            analise_limpa = analise.replace("**", "").replace("#", "").strip()

            # Construindo os placeholders manualmente para garantir a sintaxe correta
            # Queremos que o placeholder no DOCX seja {{NOTA_C1}}, {{ANALISE_C1}}, etc.
            placeholder_nota = "{{NOTA_C" + str(i) + "}}"
            substituicoes[placeholder_nota] = nota

            placeholder_analise = "{{ANALISE_C" + str(i) + "}}"
            substituicoes[placeholder_analise] = analise_limpa

        def aplicar_substituicao(paragrafo_ou_celula):
            """
            Substitui placeholders em um parágrafo ou célula de tabela,
            preservando a formatação dentro dos runs.
            """
            # Itera sobre os runs do parágrafo/célula para lidar com formatação
            # É necessário copiar o texto antes de modificar para evitar problemas de iteração
            runs_to_process = list(paragrafo_ou_celula.runs)

            # Se houver apenas um run, podemos substituir diretamente no texto dele
            if len(runs_to_process) == 1:
                texto_original_run = runs_to_process[0].text
                texto_modificado = texto_original_run
                for placeholder, valor in substituicoes.items():
                    if placeholder in texto_modificado:
                        texto_modificado = texto_modificado.replace(
                            placeholder, str(valor)
                        )

                # Atualiza o texto do run apenas se houve mudança
                if texto_modificado != texto_original_run:
                    runs_to_process[0].text = texto_modificado
            else:
                # Se houver múltiplos runs, a substituição é mais complexa.
                # Para simplicidade e para evitar quebrar formatação complexa,
                # vamos construir um novo texto e redefinir o texto do parágrafo.
                # Isso pode perder formatação mais fina, mas garante a substituição.

                texto_completo = ""
                # Concatena o texto de todos os runs
                for run in runs_to_process:
                    texto_completo += run.text

                # Aplica substituições no texto completo
                texto_modificado = texto_completo
                for placeholder, valor in substituicoes.items():
                    texto_modificado = texto_modificado.replace(placeholder, str(valor))

                # Limpa os runs existentes e insere o novo texto
                # (Isso pode perder formatação original se houver múltiplos runs com estilos diferentes)
                if texto_modificado != texto_completo:
                    # Remove todos os runs existentes
                    for r in list(paragrafo_ou_celula.runs):
                        r.text = ""
                    # Adiciona um novo run com o texto modificado
                    new_run = paragrafo_ou_celula.add_run(texto_modificado)
                    # Opcionalmente, pode-se tentar copiar a formatação do primeiro run, se existir
                    if runs_to_process:
                        try:
                            # Tenta copiar algumas propriedades básicas de formatação
                            new_run.font.name = runs_to_process[0].font.name
                            new_run.font.size = runs_to_process[0].font.size
                            new_run.font.bold = runs_to_process[0].font.bold
                            new_run.font.italic = runs_to_process[0].font.italic
                        except Exception as font_err:
                            logger.warning(
                                f"Não foi possível copiar formatação de fonte: {font_err}"
                            )

        # 1. Processa o Corpo do Texto
        for p in document.paragraphs:
            aplicar_substituicao(p)

        # 2. Processa as Tabelas do Corpo
        for tabela in document.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for p in celula.paragraphs:
                        aplicar_substituicao(p)

        # 3. Processa Cabeçalhos e Rodapés
        for section in document.sections:
            # Processa parágrafos no cabeçalho
            if section.header:
                for p in section.header.paragraphs:
                    aplicar_substituicao(p)
                # Processa tabelas dentro do cabeçalho
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                aplicar_substituicao(p)

            # Processa parágrafos no rodapé
            if section.footer:
                for p in section.footer.paragraphs:
                    aplicar_substituicao(p)
                # Processa tabelas dentro do rodapé
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                aplicar_substituicao(p)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        logger.info("Arquivo DOCX gerado com sucesso.")
        return buffer

    except FileNotFoundError:
        logger.critical(
            f"O arquivo de template '{caminho_template}' não foi encontrado."
        )
        return None
    except Exception as e:
        logger.error(f"Erro ao gerar o arquivo DOCX: {e}")
        return None

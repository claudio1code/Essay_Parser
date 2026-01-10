from io import BytesIO
from typing import Any, Dict, Optional
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.core.logger import get_logger
from config import Config

logger = get_logger(__name__)


def substituir_em_paragrafo(paragrafo: Paragraph, substituicoes: Dict[str, str]) -> None:
    """
    Substitui placeholders em um parágrafo, lidando com o problema
    de placeholders quebrados em múltiplos runs.
    
    Estratégia:
    1. Concatena todo o texto dos runs
    2. Verifica se há placeholders a substituir
    3. Se sim, reconstrói o parágrafo com o texto substituído
    """
    # Pega o texto completo do parágrafo
    texto_completo = ''.join(run.text for run in paragrafo.runs)
    
    # Verifica se há algum placeholder neste parágrafo
    tem_placeholder = any(placeholder in texto_completo for placeholder in substituicoes.keys())
    
    if not tem_placeholder:
        return  # Nada a fazer
    
    # Aplica todas as substituições
    texto_novo = texto_completo
    for placeholder, valor in substituicoes.items():
        if placeholder in texto_novo:
            texto_novo = texto_novo.replace(placeholder, str(valor))
    
    # Se o texto mudou, reconstrói o parágrafo
    if texto_novo != texto_completo:
        # Salva a formatação do primeiro run (se houver)
        formato_original = None
        if paragrafo.runs:
            primeiro_run = paragrafo.runs[0]
            formato_original = {
                'bold': primeiro_run.bold,
                'italic': primeiro_run.italic,
                'underline': primeiro_run.underline,
                'font_name': primeiro_run.font.name,
                'font_size': primeiro_run.font.size,
            }
        
        # Limpa todos os runs
        for run in paragrafo.runs:
            run.text = ''
        
        # Cria um novo run com o texto substituído
        novo_run = paragrafo.add_run(texto_novo)
        
        # Restaura a formatação original (se possível)
        if formato_original:
            try:
                novo_run.bold = formato_original['bold']
                novo_run.italic = formato_original['italic']
                novo_run.underline = formato_original['underline']
                if formato_original['font_name']:
                    novo_run.font.name = formato_original['font_name']
                if formato_original['font_size']:
                    novo_run.font.size = formato_original['font_size']
            except Exception as e:
                logger.warning(f"Não foi possível restaurar formatação: {e}")


def processar_tabela(tabela: Table, substituicoes: Dict[str, str]) -> None:
    """Processa todas as células de uma tabela."""
    for linha in tabela.rows:
        for celula in linha.cells:
            for paragrafo in celula.paragraphs:
                substituir_em_paragrafo(paragrafo, substituicoes)


def processar_secao(section, substituicoes: Dict[str, str]) -> None:
    """Processa cabeçalhos e rodapés de uma seção."""
    # Processa cabeçalho
    if section.header:
        for paragrafo in section.header.paragraphs:
            substituir_em_paragrafo(paragrafo, substituicoes)
        
        # Processa tabelas no cabeçalho
        for tabela in section.header.tables:
            processar_tabela(tabela, substituicoes)
    
    # Processa rodapé
    if section.footer:
        for paragrafo in section.footer.paragraphs:
            substituir_em_paragrafo(paragrafo, substituicoes)
        
        # Processa tabelas no rodapé
        for tabela in section.footer.tables:
            processar_tabela(tabela, substituicoes)


def processar_xpath_fallback(document: Document, substituicoes: Dict[str, str]) -> None:
    """
    Usa XPath como fallback para pegar elementos que não foram
    capturados pelas abordagens anteriores (ex: caixas de texto).
    """
    try:
        for element in document._element.xpath('.//w:t'):
            texto_original = element.text or ''
            texto_novo = texto_original
            
            for placeholder, valor in substituicoes.items():
                if placeholder in texto_novo:
                    texto_novo = texto_novo.replace(placeholder, str(valor))
            
            if texto_novo != texto_original:
                element.text = texto_novo
    except Exception as e:
        logger.warning(f"Falha no processamento XPath (ignorando): {e}")


def preencher_e_gerar_docx(
    dados: Dict[str, Any], caminho_template: str = Config.TEMPLATE_DOCX_PATH
) -> Optional[BytesIO]:
    """
    Preenche o template .docx com os dados da correção.
    
    Processa de forma robusta:
    - Corpo do documento
    - Tabelas no corpo
    - Cabeçalhos (header)
    - Rodapés (footer)
    - Caixas de texto (via XPath)
    """
    try:
        logger.info(f"📄 Abrindo template: {caminho_template}")
        document = Document(caminho_template)
        comps = dados.get("analise_competencias", {})

        # 1. Prepara o Dicionário de Substituição
        substituicoes = {
            "{{NOME_ALUNO}}": dados.get("nome_aluno", "Não identificado"),
            "{{TEMA}}": dados.get("tema_redacao", "Não identificado"),
            "{{ANO}}": dados.get("ano_turma", "Não informado"),
            "{{BIMESTRE}}": dados.get("bimestre", "Não informado"),
            "{{NOTA_FINAL}}": str(dados.get("nota_final", 0)),
            "{{COMENTARIOS}}": dados.get("comentarios_gerais", "Sem comentários."),
            "{{ALERTA_ORIGINALIDADE}}": dados.get("alerta_originalidade") or ""
        }

        # Adiciona notas e análises das competências
        for i in range(1, 6):
            comp_data = comps.get(f"c{i}", {})
            nota = str(comp_data.get("nota", 0))
            analise = comp_data.get("analise", "Análise não disponível.")
            
            # Remove markdown da análise
            analise_limpa = analise.replace("**", "").replace("#", "").strip()
            
            substituicoes[f"{{{{NOTA_C{i}}}}}"] = nota
            substituicoes[f"{{{{ANALISE_C{i}}}}}"] = analise_limpa

        logger.info("🔄 Iniciando substituições...")
        logger.info(f"   Total de placeholders: {len(substituicoes)}")

        # 2. Processa o CORPO do documento
        logger.info("   📝 Processando corpo do documento...")
        for paragrafo in document.paragraphs:
            substituir_em_paragrafo(paragrafo, substituicoes)

        # 3. Processa TABELAS no corpo
        logger.info("   📊 Processando tabelas no corpo...")
        for tabela in document.tables:
            processar_tabela(tabela, substituicoes)

        # 4. Processa CABEÇALHOS e RODAPÉS de todas as seções
        logger.info("   📋 Processando cabeçalhos e rodapés...")
        for i, section in enumerate(document.sections):
            processar_secao(section, substituicoes)

        # 5. XPath como FALLBACK (caixas de texto, elementos especiais)
        logger.info("   🔍 Processando elementos especiais (XPath)...")
        processar_xpath_fallback(document, substituicoes)

        # 6. Salva o documento
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        logger.info(f"✅ Relatório gerado com sucesso para: {dados.get('nome_aluno')}")
        return buffer

    except FileNotFoundError:
        logger.critical(f"❌ Template não encontrado: {caminho_template}")
        return None
    except Exception as e:
        logger.error(f"❌ Erro ao gerar DOCX: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return None
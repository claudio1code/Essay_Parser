import os
import sys
import argparse
from docx import Document

def analisar_documento(file_path: str):
    """
    Este script analisa um arquivo .docx para diagnosticar como os placeholders
    estão estruturados internamente, verificando se estão divididos em
    múltiplos 'runs'.
    """
    try:
        document = Document(file_path)
        print(f"✅ Analisando o arquivo: {file_path}\n")
    except Exception as e:
        print(f"❌ Erro ao abrir o arquivo: {e}")
        return

    # Lista expandida para verificar todos os placeholders
    placeholders_de_interesse = [
        "{{NOME_ALUNO}}", "{{TEMA}}", "{{ANO}}", "{{BIMESTRE}}",
        "{{NOTA_FINAL}}", "{{COMENTARIOS}}", "{{ALERTA_ORIGINALIDADE}}",
        "{{NOTA_C1}}", "{{ANALISE_C1}}",
        "{{NOTA_C2}}", "{{ANALISE_C2}}",
        "{{NOTA_C3}}", "{{ANALISE_C3}}",
        "{{NOTA_C4}}", "{{ANALISE_C4}}",
        "{{NOTA_C5}}", "{{ANALISE_C5}}",
    ]

    # Coleta todos os parágrafos do documento (corpo, tabelas, cabeçalhos)
    todos_paragrafos = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                todos_paragrafos.extend(cell.paragraphs)
    for section in document.sections:
        for p in section.header.paragraphs:
            todos_paragrafos.append(p)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    todos_paragrafos.extend(cell.paragraphs)

    print("--- INICIANDO ANÁLISE ESTRUTURAL DO DOCUMENTO ---")
    print("Analisando o conteúdo e a estrutura dos 'runs' para cada placeholder.\n")

    found_count = 0
    for p in todos_paragrafos:
        # Verifica se o texto do parágrafo contém algum placeholder de interesse
        if any(ph in p.text for ph in placeholders_de_interesse):
            found_count += 1
            print("="*70)
            print(f"📄 Texto completo do parágrafo: '{p.text.strip()}'")
            print(f"🛠️  Estrutura interna ('Runs'):")
            if len(p.runs) == 1:
                print("  -> O parágrafo inteiro está em um único 'Run'.")
            else:
                print(f"  -> O parágrafo está dividido em {len(p.runs)} 'Runs'.")
            
            for i, run in enumerate(p.runs):
                print(f"    - Run {i}: '{run.text}'")
            print("="*70 + "\n")

    if found_count == 0:
        print("\n--- RESULTADO ---")
        print("❌ Nenhum parágrafo com placeholders de interesse foi encontrado no documento.")
    else:
        print("\n--- RESULTADO ---")
        print(f"✅ Análise concluída. {found_count} parágrafos contendo placeholders foram analisados acima.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Analisa a estrutura interna de um arquivo .docx.")
    parser.add_argument("file_path", type=str, help="Caminho para o arquivo .docx a ser analisado.")
    args = parser.parse_args()
    
    analisar_documento(args.file_path)